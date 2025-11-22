from docx import Document
from pathlib import Path

OUT = Path("tests/resumes")
OUT.mkdir(parents=True, exist_ok=True)

resumes = [
    {
        "filename": "resume_1.docx",
        "name": "Alice Johnson",
        "summary": "Software engineer with experience in Python and Django.",
        "skills": ["Python", "Django", "PostgreSQL", "Docker", "AWS"],
        "education": "Bachelor of Science in Computer Science, University of Somewhere 2018",
        "experience": [
            "Software Engineer, Acme Corp (2019 - Present)\n- Worked on backend services using Python and Django.",
        ]
    },
    {
        "filename": "resume_2.docx",
        "name": "Bob Smith",
        "summary": "Data scientist specialized in machine learning and NLP.",
        "skills": ["Python", "pandas", "scikit-learn", "nlp", "TensorFlow"],
        "education": "Master of Science in Data Science, Data University 2020",
        "experience": [
            "Data Scientist, DataWorks (2020 - Present)\n- Built ML models for customer churn prediction.",
        ]
    },
    {
        "filename": "resume_3.docx",
        "name": "Carol Lee",
        "summary": "Frontend developer with React and CSS expertise.",
        "skills": ["HTML", "CSS", "React", "JavaScript", "TypeScript"],
        "education": "Bachelor of Arts in Design, Creative College 2016",
        "experience": [
            "Frontend Developer, WebStudio (2017 - 2022)\n- Led UI development using React."
        ]
    },
    {
        "filename": "resume_4.docx",
        "name": "Daniel Kim",
        "summary": "DevOps engineer focusing on CI/CD, Docker and Kubernetes.",
        "skills": ["Docker", "Kubernetes", "Jenkins", "AWS", "Terraform"],
        "education": "Bachelor of Engineering, Tech Institute 2015",
        "experience": [
            "DevOps Engineer, InfraCorp (2016 - Present)\n- Managed infrastructure and CI/CD pipelines."
        ]
    },
    {
        "filename": "resume_5.docx",
        "name": "Eva Martinez",
        "summary": "Full-stack engineer with Node.js and React experience.",
        "skills": ["Node.js", "Express", "React", "MongoDB", "GraphQL"],
        "education": "Bachelor of Science in Software Engineering, State University 2019",
        "experience": [
            "Full Stack Developer, StartupX (2019 - Present)\n- Built REST and GraphQL APIs."
        ]
    }
]

for r in resumes:
    doc = Document()
    doc.add_heading(r["name"], level=1)
    doc.add_paragraph(r["summary"])
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(r["skills"]))
    doc.add_heading("Education", level=2)
    doc.add_paragraph(r["education"])
    doc.add_heading("Experience", level=2)
    for exp in r["experience"]:
        doc.add_paragraph(exp)
    
    out_path = OUT / r["filename"]
    doc.save(out_path)

print(f"Generated {len(resumes)} test resumes in {OUT}")